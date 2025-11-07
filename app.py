from flask import Flask, render_template, request, flash, jsonify, send_file
from serpapi import GoogleSearch
from sentence_transformers import SentenceTransformer, util
from transformers import AutoTokenizer, AutoModelForSequenceClassification, GPT2TokenizerFast, GPT2LMHeadModel
import torch
import math
import fitz
import os
import re
import json
import time
from datetime import datetime
import pandas as pd
from io import BytesIO
import tempfile
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import random
import uuid

# ========== APP CONFIG ==========
app = Flask(__name__)
app.secret_key = "secret123"
os.makedirs("uploads", exist_ok=True)
os.makedirs("reports", exist_ok=True)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# ========== MODELS ==========

# BERT model for plagiarism
bert_model = SentenceTransformer("paraphrase-mpnet-base-v2")

# RoBERTa AI detector
ROBERTA_MODEL = "roberta-large-openai-detector"
roberta_tok = AutoTokenizer.from_pretrained(ROBERTA_MODEL)
roberta_model = AutoModelForSequenceClassification.from_pretrained(ROBERTA_MODEL).to(device).eval()

# GPT-2 for perplexity
GPT2_MODEL = "gpt2"
gpt2_tok = GPT2TokenizerFast.from_pretrained(GPT2_MODEL)
gpt2_model = GPT2LMHeadModel.from_pretrained(GPT2_MODEL).to(device).eval()


# ========== AI DETECTION HELPERS ==========

def roberta_ai_prob(text: str) -> float:
    """Return RoBERTa detector probability (0–1) of AI text."""
    try:
        inputs = roberta_tok(text, return_tensors="pt", truncation=True, max_length=512).to(device)
        with torch.no_grad():
            logits = roberta_model(**inputs).logits
            probs = torch.softmax(logits, dim=-1)
        return float(probs[0][1].item())
    except:
        return 0.5  # Default value if analysis fails

def gpt2_perplexity(text: str) -> float:
    """Return GPT-2 perplexity (lower = more AI-like)."""
    try:
        enc = gpt2_tok(text, return_tensors="pt", truncation=True, max_length=1024).to(device)
        input_ids = enc["input_ids"]
        with torch.no_grad():
            outputs = gpt2_model(input_ids, labels=input_ids)
            loss = outputs.loss.item()
        return math.exp(loss)
    except:
        return 100.0  # Default value if analysis fails

def normalize_perplexity(ppl: float, low=10.0, high=200.0) -> float:
    """Map perplexity to 0–1 where 1 = AI-like."""
    if ppl <= low:
        return 1.0
    if ppl >= high:
        return 0.0
    return 1.0 - (ppl - low) / (high - low)

def ensemble_ai_score(text: str, alpha=0.6, beta=0.4) -> dict:
    """Combine RoBERTa & GPT-2 signals."""
    rprob = roberta_ai_prob(text)
    ppl = gpt2_perplexity(text)
    pn = normalize_perplexity(ppl)
    ensemble = alpha * rprob + beta * pn
    return {"roberta": rprob, "perplexity": ppl, "perp_norm": pn, "ensemble": ensemble}


# ========== PLAGIARISM HELPERS ==========

def extract_text_from_pdf(path, exclude_pages=None):
    """Extract text from PDF with optional page exclusion."""
    text = ""
    with fitz.open(path) as pdf:
        total_pages = len(pdf)
        for page_num in range(total_pages):
            if exclude_pages and str(page_num + 1) in exclude_pages:
                continue
            page = pdf[page_num]
            text += page.get_text()
    return text, total_pages

def extract_text_from_docx(path, exclude_pages=None):
    """Extract text from DOCX with optional page exclusion."""
    text = ""
    doc = Document(path)
    
    # Simple text extraction for DOCX
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    return text, 1  # DOCX doesn't have clear page boundaries

def get_file_preview(path, file_type, max_pages=5):
    """Get preview of file content for page selection."""
    preview_data = {
        "total_pages": 0,
        "previews": []
    }
    
    try:
        if file_type == "pdf":
            with fitz.open(path) as pdf:
                preview_data["total_pages"] = len(pdf)
                for page_num in range(min(max_pages, len(pdf))):
                    page = pdf[page_num]
                    text = page.get_text()
                    # Get first 200 characters as preview
                    preview = text[:200].replace('\n', ' ') + "..." if len(text) > 200 else text.replace('\n', ' ')
                    preview_data["previews"].append({
                        "page": page_num + 1,
                        "preview": preview,
                        "has_content": len(text.strip()) > 0
                    })
        
        elif file_type == "docx":
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            preview_data["total_pages"] = 1  # DOCX treated as single page
            if paragraphs:
                preview = ' '.join(paragraphs)[:200] + "..." if len(' '.join(paragraphs)) > 200 else ' '.join(paragraphs)
                preview_data["previews"].append({
                    "page": 1,
                    "preview": preview,
                    "has_content": True
                })
    
    except Exception as e:
        print(f"Error getting file preview: {e}")
    
    return preview_data

def split_sentences(text):
    return re.split(r'(?<=[.!?])\s+', text.strip())

def search_serpapi(query):
    """Search sentence on web using SerpApi."""
    try:
        search = GoogleSearch({
            "q": query,
            "api_key": "demo_key",  # Use demo key for testing
            "num": 3
        })
        results = search.get_dict()
        if "organic_results" in results:
            return results["organic_results"][:3]
    except Exception as e:
        print("Search error:", e)
    return []

def sentence_similarity(sentence, results):
    """Find plagiarism similarity using BERT embeddings."""
    try:
        sent_emb = bert_model.encode(sentence, convert_to_tensor=True)
        best_score, best_match = 0, {}
        for r in results:
            snippet = r.get("snippet", "")
            if snippet:
                sn_emb = bert_model.encode(snippet, convert_to_tensor=True)
                score = float(util.cos_sim(sent_emb, sn_emb))
                if score > best_score:
                    best_score = score
                    best_match = {
                        "title": r.get("title"),
                        "link": r.get("link"),
                        "snippet": snippet,
                        "score": round(score, 3)
                    }
        return best_match or {"score": 0}
    except:
        return {"score": 0}

def calculate_overall_risk(plagiarism_score, ai_score):
    """Calculate overall risk level."""
    if plagiarism_score > 50 or ai_score > 50:
        return "High"
    elif plagiarism_score > 25 or ai_score > 25:
        return "Medium"
    else:
        return "Low"

# ========== PDF EXPORT HELPERS ==========

def get_risk_color(risk_level):
    """Get color for risk level."""
    color_map = {
        "High": colors.red,
        "Medium": colors.orange,
        "Low": colors.green
    }
    return color_map.get(risk_level, colors.black)

def get_score_color(score):
    """Get color for score."""
    if score > 50:
        return colors.red
    elif score > 25:
        return colors.orange
    else:
        return colors.green

def create_pdf_report(analysis_data):
    """Create a PDF report from analysis data."""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              topMargin=1*inch, 
                              rightMargin=0.5*inch, 
                              leftMargin=0.5*inch, 
                              bottomMargin=1*inch)
        
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.HexColor('#0b5ed7')
        )
        title = Paragraph("AI + BERT Plagiarism Analysis Report", title_style)
        story.append(title)
        
        # Date
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=1,
            textColor=colors.gray
        )
        date = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style)
        story.append(date)
        story.append(Spacer(1, 20))
        
        # Summary Section
        summary_style = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#0b5ed7')
        )
        summary_title = Paragraph("Executive Summary", summary_style)
        story.append(summary_title)
        
        # Summary Table
        summary_data = [
            ['Metric', 'Score', 'Risk Level'],
            ['Plagiarism Score', f"{analysis_data['avg_similarity']}%", analysis_data['risk_level']],
            ['AI Probability', f"{analysis_data['avg_ai']}%", analysis_data['risk_level']],
            ['Sentences Analyzed', str(analysis_data['total_sentences']), '-'],
            ['Processing Time', f"{analysis_data['processing_time']} seconds", '-']
        ]
        
        summary_table = Table(summary_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0b5ed7')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        # Color code the scores
        plagiarism_color = get_score_color(analysis_data['avg_similarity'])
        ai_color = get_score_color(analysis_data['avg_ai'])
        risk_color = get_risk_color(analysis_data['risk_level'])
        
        summary_table.setStyle(TableStyle([
            ('TEXTCOLOR', (1, 1), (1, 1), plagiarism_color),
            ('TEXTCOLOR', (1, 2), (1, 2), ai_color),
            ('TEXTCOLOR', (2, 1), (2, 2), risk_color),
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # Detailed Analysis Section
        detailed_title = Paragraph("Detailed Analysis", summary_style)
        story.append(detailed_title)
        
        # Risk Breakdown
        high_plagiarism_pct = (analysis_data['high_plagiarism_count'] / analysis_data['total_sentences'] * 100) if analysis_data['total_sentences'] > 0 else 0
        high_ai_pct = (analysis_data['high_ai_count'] / analysis_data['total_sentences'] * 100) if analysis_data['total_sentences'] > 0 else 0
        
        risk_data = [
            ['Risk Category', 'Count', 'Percentage'],
            ['High Plagiarism', str(analysis_data['high_plagiarism_count']), f"{high_plagiarism_pct:.1f}%"],
            ['High AI Probability', str(analysis_data['high_ai_count']), f"{high_ai_pct:.1f}%"]
        ]
        
        risk_table = Table(risk_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6c757d')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(risk_table)
        story.append(Spacer(1, 20))
        
        # Sentence Analysis Section (limited to first 5 sentences)
        if analysis_data.get('highlights') and len(analysis_data['highlights']) > 0:
            sentences_title = Paragraph("Sentence Analysis Highlights", summary_style)
            story.append(sentences_title)
            
            # Add top 5 highlighted sentences
            sentence_data = [['Sentence', 'Similarity', 'AI Prob', 'Status']]
            for i, highlight in enumerate(analysis_data['highlights'][:5]):  # Limit to first 5
                status = "High Risk" if highlight['color'] == 'red' else "Medium Risk" if highlight['color'] == 'orange' else "Low Risk"
                sentence_text = highlight['sentence'][:80] + "..." if len(highlight['sentence']) > 80 else highlight['sentence']
                sentence_data.append([
                    Paragraph(sentence_text, styles['Normal']),
                    f"{highlight['score']}%",
                    f"{highlight['ai']}%",
                    status
                ])
            
            sentence_table = Table(sentence_data, colWidths=[3*inch, 1*inch, 1*inch, 1.5*inch])
            sentence_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0b5ed7')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(sentence_table)
            story.append(Spacer(1, 20))
        
        # Interpretation Guide
        guide_title = Paragraph("Interpretation Guide", summary_style)
        story.append(guide_title)
        
        guide_data = [
            ['Color', 'Similarity Range', 'Risk Level', 'Description'],
            ['Green', '0-50%', 'Low', 'Minimal plagiarism risk'],
            ['Orange', '50-75%', 'Medium', 'Moderate plagiarism risk - review recommended'],
            ['Red', '75-100%', 'High', 'High plagiarism risk - immediate attention required']
        ]
        
        guide_table = Table(guide_data, colWidths=[1*inch, 1.5*inch, 1.5*inch, 2.5*inch])
        guide_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6c757d')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 1), (0, 1), colors.green),
            ('BACKGROUND', (0, 2), (0, 2), colors.orange),
            ('BACKGROUND', (0, 3), (0, 3), colors.red),
            ('TEXTCOLOR', (0, 1), (0, 3), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(guide_table)
        story.append(Spacer(1, 20))
        
        # Footer
        footer_style = ParagraphStyle(
            'FooterStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=1,
            textColor=colors.gray
        )
        footer = Paragraph("Generated by AI + BERT Plagiarism Checker | Confidential Report", footer_style)
        story.append(footer)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Error creating PDF: {e}")
        raise e

# ========== ROUTES ==========

@app.route("/")
def index():
    return render_template("index.html")

# Store uploaded files temporarily
uploaded_files = {}

@app.route("/upload", methods=["POST"])
def handle_upload():
    """Handle file upload and return preview for page selection."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400
    
    # Save uploaded file temporarily
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, file.filename)
    file.save(file_path)
    
    # Determine file type
    file_type = None
    if file.filename.lower().endswith('.pdf'):
        file_type = "pdf"
    elif file.filename.lower().endswith(('.doc', '.docx')):
        file_type = "docx"
    else:
        os.remove(file_path)
        return jsonify({"error": "Unsupported file type. Please upload PDF or Word document."}), 400
    
    # Get file preview
    preview_data = get_file_preview(file_path, file_type)
    preview_data["filename"] = file.filename
    preview_data["file_type"] = file_type
    
    # Store file info
    file_id = str(uuid.uuid4())
    uploaded_files[file_id] = {
        "path": file_path,
        "type": file_type,
        "filename": file.filename
    }
    
    preview_data["file_id"] = file_id
    
    return jsonify(preview_data)

@app.route("/check", methods=["POST"])
def check():
    start_time = time.time()
    
    text = ""
    total_pages = 0
    deep_search = request.form.get("deep_search") == "on"
    exclude_pages = request.form.get("exclude_pages", "").split(",") if request.form.get("exclude_pages") else []
    file_id = request.form.get("file_id")
    
    # Check if we have a file from the upload step
    if file_id and file_id in uploaded_files:
        file_info = uploaded_files[file_id]
        file_path = file_info["path"]
        file_type = file_info["type"]
        
        if os.path.exists(file_path):
            if file_type == "pdf":
                text, total_pages = extract_text_from_pdf(file_path, exclude_pages)
            elif file_type == "docx":
                text, total_pages = extract_text_from_docx(file_path, exclude_pages)
            
            # Clean up temporary file
            try:
                os.remove(file_path)
                del uploaded_files[file_id]
            except:
                pass
    else:
        # Handle direct text input
        text = request.form.get("text", "")

    if not text.strip():
        flash("Please paste text or upload a PDF/Word document.")
        return render_template("index.html", original_text=request.form.get("text", ""))

    sentences = [s for s in split_sentences(text) if len(s) > 20]
    if not sentences:
        flash("Text too short for analysis.")
        return render_template("index.html", original_text=text)

    results = []
    high_plagiarism_count = 0
    high_ai_count = 0

    # Use demo mode for testing
    demo_mode = True
    
    for i, s in enumerate(sentences):
        if demo_mode:
            # Generate demo data for testing
            sim_score = round(random.uniform(0.1, 0.9), 2)
            ai_score = round(random.uniform(10, 90), 2)
            
            match = {
                "score": sim_score,
                "title": "Sample Source" if sim_score > 0.5 else None,
                "link": "https://example.com" if sim_score > 0.5 else None,
                "snippet": "This is a sample matching text snippet..." if sim_score > 0.5 else None
            }
            
            ai_details = {
                "roberta": round(random.uniform(0.1, 0.9), 2),
                "perplexity": round(random.uniform(20, 150), 2),
                "perp_norm": round(random.uniform(0.1, 0.9), 2),
                "ensemble": ai_score / 100
            }
        else:
            # Perform actual analysis
            serp = search_serpapi(s) if deep_search else []
            match = sentence_similarity(s, serp)
            ai_details = ensemble_ai_score(s)
            ai_score = round(ai_details["ensemble"] * 100, 2)

        # Count high risk sentences
        if match.get("score", 0) > 0.75:
            high_plagiarism_count += 1
        if ai_score > 50:
            high_ai_count += 1

        results.append({
            "sentence": s,
            **match,
            "ai_score": ai_score,
            "ai_details": ai_details
        })

    # Highlighting
    highlights = []
    for r in results:
        sim = r.get("score", 0)
        color = "green"
        if sim > 0.75:
            color = "red"
        elif sim > 0.5:
            color = "orange"
        highlights.append({
            "sentence": r["sentence"],
            "color": color,
            "score": int(sim * 100),
            "ai": r["ai_score"]
        })

    avg_sim = round(sum([r.get("score", 0) for r in results]) / len(results) * 100, 2)
    avg_ai = round(sum([r["ai_score"] for r in results]) / len(results), 2)
    
    # Calculate processing time
    processing_time = round(time.time() - start_time, 2)
    
    # Calculate risk level
    risk_level = calculate_overall_risk(avg_sim, avg_ai)

    chart_labels = [f"S{i+1}" for i in range(len(results))]
    chart_scores = [r["ai_score"] for r in results]

    # Prepare data for JavaScript
    analysis_data = {
        "avg_similarity": avg_sim,
        "avg_ai": avg_ai,
        "total_sentences": len(sentences),
        "risk_level": risk_level,
        "processing_time": processing_time,
        "high_plagiarism_count": high_plagiarism_count,
        "high_ai_count": high_ai_count,
        "chart_labels": chart_labels,
        "chart_scores": chart_scores,
        "highlights": highlights,
        "sentence_results": results,
        "total_pages": total_pages,
        "excluded_pages": exclude_pages
    }

    return render_template(
        "index.html",
        original_text=text,
        analysis_data=json.dumps(analysis_data)
    )

@app.route("/export-pdf", methods=["POST"])
def export_pdf():
    """Export analysis report as PDF."""
    try:
        analysis_data = request.get_json()
        
        if not analysis_data:
            return jsonify({"error": "No analysis data provided"}), 400
        
        # Create PDF
        pdf_buffer = create_pdf_report(analysis_data)
        
        # Return PDF as download
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=f"plagiarism_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"PDF export error: {e}")
        return jsonify({"error": f"Failed to generate PDF: {str(e)}"}), 500

@app.route("/demo", methods=["POST"])
def demo_analysis():
    """Demo endpoint for testing without actual analysis"""
    text = request.form.get("text", "")
    
    if not text.strip():
        return jsonify({"error": "No text provided"}), 400
    
    sentences = [s for s in split_sentences(text) if len(s) > 20]
    
    if not sentences:
        return jsonify({"error": "Text too short"}), 400
    
    # Generate demo data
    results = []
    highlights = []
    
    for i, s in enumerate(sentences):
        sim_score = round(random.uniform(0.1, 0.9), 2)
        ai_score = round(random.uniform(10, 90), 2)
        
        color = "green"
        if sim_score > 0.75:
            color = "red"
        elif sim_score > 0.5:
            color = "orange"
            
        results.append({
            "sentence": s,
            "score": sim_score,
            "ai_score": ai_score,
            "title": "Sample Source" if sim_score > 0.5 else None
        })
        
        highlights.append({
            "sentence": s,
            "color": color,
            "score": int(sim_score * 100),
            "ai": ai_score
        })
    
    avg_sim = round(sum([r["score"] for r in results]) / len(results) * 100, 2)
    avg_ai = round(sum([r["ai_score"] for r in results]) / len(results), 2)
    
    analysis_data = {
        "avg_similarity": avg_sim,
        "avg_ai": avg_ai,
        "total_sentences": len(sentences),
        "risk_level": calculate_overall_risk(avg_sim, avg_ai),
        "processing_time": round(random.uniform(1, 5), 2),
        "high_plagiarism_count": sum(1 for r in results if r["score"] > 0.75),
        "high_ai_count": sum(1 for r in results if r["ai_score"] > 50),
        "chart_labels": [f"S{i+1}" for i in range(len(results))],
        "chart_scores": [r["ai_score"] for r in results],
        "highlights": highlights,
        "sentence_results": results,
        "total_pages": 0,
        "excluded_pages": []
    }
    
    return jsonify(analysis_data)

if __name__ == "__main__":
    app.run(debug=True, port=5000)